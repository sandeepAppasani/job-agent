"""Utilities for reading resumes and managing versioned application folders."""
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
import pdfplumber

from utils.logger import get_logger

logger = get_logger(__name__)


def read_resume_text(resume_path: Path) -> str:
    """Extract plain text from a .docx or .pdf resume."""
    suffix = resume_path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(resume_path)
    elif suffix == ".pdf":
        return _read_pdf(resume_path)
    else:
        raise ValueError(f"Unsupported resume format: {suffix}")


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    with pdfplumber.open(str(path)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def create_application_folder(
    base_dir: Path,
    job_title: str,
    company: str,
    description_snippet: str,
) -> Path:
    """
    Create a versioned folder:  applications/YYYYMMDD_JobTitle_Company_Snippet/
    Returns the created folder path.
    """
    date_str = datetime.now().strftime("%Y%m%d")
    safe = re.compile(r"[^a-zA-Z0-9 _-]")
    title_slug = safe.sub("", job_title).strip().replace(" ", "_")[:30]
    company_slug = safe.sub("", company).strip().replace(" ", "_")[:20]
    snippet_slug = safe.sub("", description_snippet).strip().replace(" ", "_")[:20]

    folder_name = f"{date_str}_{title_slug}_{company_slug}_{snippet_slug}"
    folder = base_dir / folder_name
    folder.mkdir(parents=True, exist_ok=True)
    logger.info(f"Created application folder: {folder}")
    return folder


def save_tailored_resume(
    original_resume_path: Path,
    tailored_text: str,
    output_folder: Path,
    filename_prefix: str = "tailored_resume",
) -> Path:
    """
    Save a tailored resume as a .docx file in the output folder.
    Also copies the original resume alongside it for reference.
    Returns the path of the tailored .docx.
    """
    # Copy original resume
    original_copy = output_folder / f"original_{original_resume_path.name}"
    shutil.copy2(original_resume_path, original_copy)

    # Save tailored version as .docx
    doc = Document()
    for line in tailored_text.splitlines():
        doc.add_paragraph(line)

    output_path = output_folder / f"{filename_prefix}.docx"
    doc.save(str(output_path))
    logger.info(f"Saved tailored resume: {output_path}")
    return output_path


def get_resume_version(folder: Path) -> str:
    """Return a short version string based on the folder name."""
    return folder.name
